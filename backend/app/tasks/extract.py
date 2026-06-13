"""Celery task: extract text from RFP document and run AI requirement extraction."""

import asyncio
import os
from pathlib import Path
from typing import Optional
from uuid import UUID

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from app.config import settings
from app.db.models import DocumentStatus, JobStatus
from app.logger import get_logger
from app.tasks.celery_app import celery_app

logger = get_logger("task_extract")


def _run_async(coro):
    """Run an async coroutine in a Celery worker context."""
    loop = asyncio.new_event_loop()
    try:
        return loop.run_until_complete(coro)
    finally:
        loop.close()


@celery_app.task(
    name="tasks.extract.extract_rfp_text",
    bind=True,
    max_retries=3,
    default_retry_delay=30,
)
def extract_rfp_text(self, job_id: str, document_id: str) -> dict:
    """Extract text from an RFP document and run AI requirement extraction.

    Steps:
        1. Load document from storage
        2. Extract text with PyMuPDF (PDF) or python-docx (DOCX)
        3. Save extracted text to DB
        4. Run AI extraction to identify requirements
        5. Save requirements to DB
        6. Update job to DONE

    Args:
        job_id: UUID string of the Job record.
        document_id: UUID string of the RfpDocument record.

    Returns:
        Dict with n_requirements and status.
    """
    return _run_async(_extract_rfp_text_async(self, job_id, document_id))


async def _extract_rfp_text_async(task, job_id: str, document_id: str) -> dict:
    """Async implementation of the extraction task."""
    from app.db.database import get_db_session
    from app.db.models import Job, RfpDocument
    from app.db.repository import db_get_by_id, db_update
    from app.ai.extractor import RFPExtractor

    async with get_db_session() as session:
        job = await db_get_by_id(session, Job, UUID(job_id))
        doc = await db_get_by_id(session, RfpDocument, UUID(document_id))

        if not job or not doc:
            logger.error("Job or document not found", job_id=job_id, document_id=document_id)
            return {"status": "failed", "error": "Job or document not found"}

        try:
            # Update status to processing
            await db_update(session, Job, job.id, {"status": JobStatus.PROCESSING, "progress_pct": 5})

            # Step 1: Locate file
            file_path = _resolve_file_path(doc.s3_key)
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")

            task.update_state(state="STARTED", meta={"progress": 10})

            # Step 2: Extract raw text
            ext = Path(doc.filename).suffix.lower()
            if ext == ".pdf":
                extracted_text = _extract_pdf(file_path)
            elif ext in (".docx", ".doc"):
                extracted_text = _extract_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {ext}")

            logger.info(
                "Text extracted",
                document_id=document_id,
                char_count=len(extracted_text),
            )

            # Step 3: Save text to document
            await db_update(
                session,
                RfpDocument,
                doc.id,
                {
                    "extracted_text": extracted_text,
                    "status": DocumentStatus.PROCESSING,
                },
            )

            await db_update(session, Job, job.id, {"progress_pct": 30})
            task.update_state(state="STARTED", meta={"progress": 30})

            # Step 4: AI extraction → requirements
            extractor = RFPExtractor()

            async def progress_cb(pct: int):
                mapped = 30 + int(pct * 0.6)  # maps 0-100 → 30-90
                job.progress_pct = mapped
                session.add(job)
                await session.commit()
                task.update_state(state="STARTED", meta={"progress": mapped})

            requirements = await extractor.extract_and_save(
                document_id=UUID(document_id),
                session=session,
                progress_callback=progress_cb,
            )

            # Step 5: Queue capability matching
            from app.tasks.match import match_capabilities
            from app.db.models import Job as JobModel
            from app.db.repository import db_create

            match_job = JobModel(
                workspace_id=doc.workspace_id,
                job_type="match",
                status=JobStatus.PENDING,
                progress_pct=0,
            )
            await db_create(session, match_job)

            celery_app.send_task(
                "tasks.match.match_capabilities",
                args=[str(match_job.id), str(doc.workspace_id)],
                task_id=str(match_job.id),
            )

            # Step 6: Mark job done
            await db_update(session, Job, job.id, {"status": JobStatus.DONE, "progress_pct": 100})

            logger.info(
                "Extraction task complete",
                document_id=document_id,
                n_requirements=len(requirements),
            )
            return {"status": "done", "n_requirements": len(requirements)}

        except Exception as exc:
            logger.error("Extraction task failed", error=str(exc), job_id=job_id)
            await db_update(
                session,
                RfpDocument,
                doc.id,
                {"status": DocumentStatus.FAILED},
            )
            await db_update(
                session,
                Job,
                job.id,
                {"status": JobStatus.FAILED, "error_msg": str(exc)},
            )

            try:
                raise task.retry(exc=exc)
            except Exception:
                return {"status": "failed", "error": str(exc)}


def _resolve_file_path(s3_key: str) -> str:
    """Resolve local file path from s3_key."""
    if settings.USE_LOCAL_STORAGE:
        return os.path.join(settings.LOCAL_STORAGE_PATH, *s3_key.split("/")[-3:])
    # For real S3, would download to temp file
    raise NotImplementedError("S3 download not yet implemented in this task")


def _extract_pdf(file_path: str) -> str:
    """Extract text from a PDF using PyMuPDF."""
    text_parts = []
    with fitz.open(file_path) as pdf:
        for page_num in range(len(pdf)):
            page = pdf[page_num]
            text_parts.append(page.get_text("text"))
    full_text = "\n\n".join(text_parts)
    logger.debug("PDF extracted", pages=len(text_parts), chars=len(full_text))
    return full_text


def _extract_docx(file_path: str) -> str:
    """Extract text from a DOCX using python-docx."""
    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Include table content
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    full_text = "\n\n".join(paragraphs)
    logger.debug("DOCX extracted", paragraphs=len(paragraphs), chars=len(full_text))
    return full_text
