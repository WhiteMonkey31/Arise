"""DOCX export endpoint for full proposal document generation."""

import io
from datetime import datetime
from typing import Optional
from uuid import UUID

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlmodel import select

from app.auth.users import current_active_user
from app.db.database import get_db
from app.db.models import (
    Capability,
    ComplianceItem,
    ComplianceStatus,
    Proposal,
    ProposalStatus,
    Requirement,
    User,
    Workspace,
)
from app.db.repository import DbSession
from app.logger import get_logger

logger = get_logger("export_api")
router = APIRouter(prefix="/api/workspaces", tags=["export"])


# ---------------------------------------------------------------------------
# Endpoint
# ---------------------------------------------------------------------------

@router.post("/{workspace_id}/export")
async def export_proposal(
    workspace_id: UUID,
    session: DbSession = Depends(get_db),
    current_user: User = Depends(current_active_user),
) -> StreamingResponse:
    """Generate and download a full DOCX proposal document.

    The document includes:
    - Cover page with workspace/org details
    - Table of contents
    - Compliance summary table
    - All proposal sections with evidence references
    - Appendix with capability details
    """
    workspace = await _check_workspace(workspace_id, current_user, session)

    # Load proposals
    from app.db.repository import db_query
    proposals = await db_query(
        session,
        Proposal,
        filters=[("workspace_id", "==", workspace_id)],
        order_by=("created_at", "asc"),
    )

    # Load compliance items
    compliance_items = await db_query(
        session,
        ComplianceItem,
        filters=[("workspace_id", "==", workspace_id)],
    )

    # Generate DOCX
    doc_bytes = await _build_docx(workspace, proposals, compliance_items, session)

    filename = f"proposal_{workspace.name.replace(' ', '_')}_{datetime.utcnow().strftime('%Y%m%d')}.docx"

    logger.info("DOCX export generated", workspace_id=str(workspace_id), filename=filename)

    return StreamingResponse(
        io.BytesIO(doc_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# DOCX Builder
# ---------------------------------------------------------------------------

async def _build_docx(
    workspace: Workspace,
    proposals: list,
    compliance_items: list,
    session: DbSession,
) -> bytes:
    """Build the full DOCX document and return as bytes."""
    doc = DocxDocument()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # ---- Cover Page ----
    _add_cover_page(doc, workspace)

    # ---- Page break ----
    doc.add_page_break()

    # ---- Table of Contents ----
    _add_toc(doc, proposals)
    doc.add_page_break()

    # ---- Compliance Summary ----
    heading = doc.add_heading("Compliance Summary", level=1)
    _add_compliance_table(doc, compliance_items, session)
    doc.add_page_break()

    # ---- Proposal Sections ----
    doc.add_heading("Proposal Sections", level=1)
    for proposal in proposals:
        await _add_proposal_section(doc, proposal, compliance_items, session)

    # ---- Appendix ----
    doc.add_page_break()
    doc.add_heading("Appendix: Capability Evidence", level=1)
    await _add_appendix(doc, compliance_items, session)

    # Serialize to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _add_cover_page(doc: DocxDocument, workspace: Workspace) -> None:
    """Add a styled cover page."""
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_heading(workspace.name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    subtitle = doc.add_paragraph("Proposal Response Document")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)

    doc.add_paragraph()
    info_para = doc.add_paragraph(
        f"Sector: {workspace.sector or 'N/A'}\n"
        f"Status: {workspace.status.value}\n"
        f"Deadline: {workspace.deadline.strftime('%B %d, %Y') if workspace.deadline else 'TBD'}\n"
        f"Prepared: {datetime.utcnow().strftime('%B %d, %Y')}"
    )
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_toc(doc: DocxDocument, proposals: list) -> None:
    """Add a simple table of contents."""
    doc.add_heading("Table of Contents", level=1)
    for i, proposal in enumerate(proposals, 1):
        toc_para = doc.add_paragraph(f"{i}. {proposal.section_title}", style="List Number")


def _add_compliance_table(
    doc: DocxDocument,
    compliance_items: list,
    session,
) -> None:
    """Add a compliance summary table."""
    if not compliance_items:
        doc.add_paragraph("No compliance data available.")
        return

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # Header row
    headers = ["Requirement", "Status", "Match Score", "Notes"]
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        run = header_cells[i].paragraphs[0].runs[0]
        run.bold = True

    # Data rows (sync helper — items already loaded)
    for item in compliance_items:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.requirement_id)[:8] + "..."
        row_cells[1].text = item.status.value
        row_cells[2].text = f"{item.match_score:.2f}" if item.match_score else "N/A"
        row_cells[3].text = item.notes or ""

        # Colour status cell
        status_run = row_cells[1].paragraphs[0].runs
        if item.status == ComplianceStatus.PASS:
            if status_run:
                status_run[0].font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        elif item.status == ComplianceStatus.GAP:
            if status_run:
                status_run[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


async def _add_proposal_section(
    doc: DocxDocument,
    proposal: Proposal,
    compliance_items: list,
    session: DbSession,
) -> None:
    """Add a single proposal section to the document."""
    doc.add_heading(proposal.section_title, level=2)

    content = proposal.current_content or proposal.ai_draft or ""
    if content:
        for paragraph_text in content.split("\n\n"):
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text.strip())
    else:
        doc.add_paragraph("(No content drafted)")

    # Quality badge
    if proposal.quality_badge:
        badge_para = doc.add_paragraph(f"Quality: {proposal.quality_badge}")
        badge_para.runs[0].bold = True
        badge_para.runs[0].font.size = Pt(9)

    doc.add_paragraph()


async def _add_appendix(
    doc: DocxDocument,
    compliance_items: list,
    session: DbSession,
) -> None:
    """Add appendix with capability details."""
    from app.db.repository import db_get_by_id
    seen_cap_ids = set()
    for item in compliance_items:
        if item.capability_id and item.capability_id not in seen_cap_ids:
            seen_cap_ids.add(item.capability_id)
            cap = await db_get_by_id(session, Capability, item.capability_id)
            if cap:
                doc.add_heading(cap.title, level=3)
                doc.add_paragraph(f"Domain: {cap.domain or 'N/A'} | Certification: {cap.certification or 'N/A'} | Year: {cap.year or 'N/A'}")
                doc.add_paragraph(cap.description)
                doc.add_paragraph()


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

async def _check_workspace(workspace_id: UUID, user: User, session: DbSession) -> Workspace:
    from app.db.repository import db_get_by_id
    workspace = await db_get_by_id(session, Workspace, workspace_id)
    if not workspace or workspace.deleted_at is not None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    if workspace.org_id != user.org_id:
        raise HTTPException(status_code=403, detail="Access denied")
    return workspace
